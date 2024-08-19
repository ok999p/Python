import pandas as pd
from docx import Document

# อ่านข้อมูลจากไฟล์ Excel
excel_file = r"C:\Users\user\Downloads\student_data_with_images2 (1).xlsx"
sheets = ['64', '65', '66', '67']

# สร้างเอกสาร Word ใหม่
doc = Document()

# เขียนหัวข้อหลักในเอกสาร Word
doc.add_heading('คณะกรรมการ การเลือกตั้ง มหาวิทยาลัยวลัยลักษณ์', level=1)
doc.add_heading('ผลคะแนนการเลือกตั้งสภานักศึกษา(ซ่อม) ประจำปีการศึกษา 2567', level=2)

for sheet in sheets:
    # อ่านข้อมูลจากแต่ละชีท
    df = pd.read_excel(excel_file, sheet_name=sheet)

    # เขียนหัวข้อรหัสในเอกสาร
    doc.add_heading(f'รหัส {sheet}', level=3)

    # ดึงข้อมูลเฉพาะที่ต้องการ
    for index, row in df.iterrows():
        student_code = row['รหัสนักศึกษา']
        name = row['ชื่อ-สกุล']
        point = row['Point']  # ปรับเป็นตัวใหญ่ 'Point'
        order = index + 1

        # เขียนข้อมูลในรูปแบบ "ลำดับที่ (รหัสนักศึกษา) (ชื่อ-สกุล) (Point)"
        doc.add_paragraph(f'ลำดับที่ {order}    {student_code}  {name}      \t\tคะแนนเสียง {point}')

# บันทึกไฟล์ Word
try:
    doc.save(r'C:\Users\user\OneDrive - Walailak University\Desktop\ผลการเลือกตั้งสภานักศึกษา.docx')
    print("ไฟล์ถูกบันทึกเรียบร้อยแล้ว")
except Exception as e:
    print(f"เกิดข้อผิดพลาดในการบันทึกไฟล์: {e}")

